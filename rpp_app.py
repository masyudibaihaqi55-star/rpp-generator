import streamlit as st
from openai import OpenAI
from docx import Document
from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

# --- SETUP OPENAI ---
import os
api_key = os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=api_key)

# --- Judul Aplikasi ---
st.title("📘 Generator RPP Otomatis")
st.write("Isi data berikut, lalu AI akan membuat RPP + LKPD + Rubrik lengkap siap diunduh dalam format Word atau PDF.")

# --- Form Input ---
with st.form("rpp_form"):
    sekolah = st.text_input("Nama Sekolah", "SMP Negeri 1")
    mapel = st.text_input("Mata Pelajaran", "PPKn")
    kelas = st.text_input("Kelas", "VII")
    semester = st.text_input("Semester", "1")
    materi = st.text_area("Materi", "Norma serta hak dan kewajiban dalam kehidupan berbangsa dan bernegara")
    topik = st.text_input("Topik", "Hak dan kewajiban anak")
    cp = st.text_area("Capaian Pembelajaran", "Memahami pentingnya norma, hak, dan kewajiban dalam kehidupan sehari-hari")
    tp = st.text_area("Tujuan Pembelajaran", "Siswa mampu mengidentifikasi contoh hak dan kewajiban anak dalam kehidupan sehari-hari")

    submit = st.form_submit_button("🚀 Buat RPP")

# --- Prompt Master ---
prompt_master = f"""
Buatkan saya RPP Pembelajaran Mendalam dengan format naratif berikut.

*Identitas RPP*
- Sekolah: {sekolah}
- Mata Pelajaran: {mapel}
- Kelas: {kelas}
- Semester: {semester}
- Fase: [[AI tentukan otomatis]]
- Alokasi Waktu: 2 JP

1. Identifikasi Peserta Didik → uraikan karakteristik umum siswa sesuai jenjang.

2. Materi Pelajaran: {materi}

3. Topik Pembelajaran: {topik}

4. 8 Dimensi Profil Lulusan → pilih sesuai relevansi.

5. Desain Pembelajaran:
   - CP: {cp}
   - TP: {tp}
   - Model Pembelajaran: AI pilih otomatis (Discovery, Inquiry, PBL, PjBL sesuai TP)
   - Sintaks: gunakan tahapan sesuai model pembelajaran terpilih.
   - Lingkungan Belajar: sebutkan ruang nyata yang digunakan.
   - Pemanfaatan Digital: uraikan media digital di awal, inti, penutup.

6. Langkah-langkah Pembelajaran (naratif seperti skenario, bukan poin singkat):
   - Kegiatan Awal → salam, doa, ice breaking.
   - Kegiatan Inti → naratif sesuai sintaks model.
   - Kegiatan Penutup → ringkasan, refleksi, apresiasi, doa.

7. Asesmen Pembelajaran (naratif, formatif & sumatif).

8. LKPD (naratif + tabel rubrik 1–4 dengan deskripsi jelas).
   - Minimal 2 LKPD berbasis sintaks.
   - Tambahkan LKPD Refleksi.
   - Gunakan tabel untuk rubrik (level 1–4).

9. Refleksi & RTL guru-siswa.

Output berupa teks naratif lengkap, rapi, dan langsung bisa digunakan sebagai RPP.
"""

# --- Proses ke OpenAI ---
if submit:
    with st.spinner("AI sedang menyusun RPP..."):
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "Kamu adalah asisten pendidikan yang ahli menyusun RPP."},
                {"role": "user", "content": prompt_master}
            ],
            temperature=0.7
        )
        hasil_rpp = response.choices[0].message.content

    # --- Tampilkan Hasil ---
    st.subheader("📖 Hasil RPP")
    st.write(hasil_rpp)

    # --- Buat Word File ---
    doc = Document()
    doc.add_heading("Rencana Pelaksanaan Pembelajaran (RPP)", level=1)
    for line in hasil_rpp.split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    buffer_word = BytesIO()
    doc.save(buffer_word)
    buffer_word.seek(0)

    # --- Buat PDF File ---
    buffer_pdf = BytesIO()
    doc_pdf = SimpleDocTemplate(buffer_pdf, pagesize=A4)
    styles = getSampleStyleSheet()
    story = [Paragraph("Rencana Pelaksanaan Pembelajaran (RPP)", styles["Heading1"]), Spacer(1, 12)]
    for line in hasil_rpp.split("\n"):
        if line.strip():
            story.append(Paragraph(line.strip(), styles["Normal"]))
            story.append(Spacer(1, 8))
    doc_pdf.build(story)
    buffer_pdf.seek(0)

    # --- Tombol Download ---
    col1, col2 = st.columns(2)
    with col1:
        st.download_button(
            "📥 Download RPP (Word)",
            data=buffer_word,
            file_name="RPP.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    with col2:
        st.download_button(
            "📥 Download RPP (PDF)",
            data=buffer_pdf,
            file_name="RPP.pdf",
            mime="application/pdf"
        )
